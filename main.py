from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import io
import re
from docx import Document
from copy import deepcopy
from docx.oxml import OxmlElement

app = FastAPI(title="Italify API Offline")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── DATABASE LOKAL (DIPATENKAN) ───
INDO_DB = set()
ENG_DB = set()

@app.on_event("startup")
def load_databases():
    global INDO_DB, ENG_DB
    print("⏳ Memuat Kamus Lokal (Paten)...")
    try:
        with open("kbbi.txt", "r", encoding="utf-8") as f:
            INDO_DB = {line.strip().lower() for line in f if line.strip()}
        print(f"✅ KBBI (Indo) siap: {len(INDO_DB)} kata")
    except FileNotFoundError:
        print("❌ File kbbi.txt tidak ditemukan!")

    try:
        with open("words_alpha.txt", "r", encoding="utf-8") as f:
            ENG_DB = {line.strip().lower() for line in f if line.strip()}
        print(f"✅ Kamus Inggris siap: {len(ENG_DB)} kata")
    except FileNotFoundError:
        print("❌ File words_alpha.txt tidak ditemukan!")


# ─── 🧠 NLP ENGINE: KATA PER KATA (SIMPLE & PRESISI) 🧠 ───

def is_valid_indo(kata):
    """Cek KBBI murni + Pengupas Imbuhan"""
    if not kata: return False
    
    # Penanganan kata ber-strip (contoh: uji-coba)
    if '-' in kata:
        return all(is_valid_indo(p) for p in kata.split('-') if p)
        
    if kata in INDO_DB: return True
    
    prefixes = ['', 'di', 'ke', 'se', 'ter', 'ber', 'me', 'mem', 'men', 'meng', 'meny', 'pe', 'pem', 'pen', 'peng', 'peny', 'diper', 'member', 'diber', 'pra', 'pasca', 'terde']
    suffixes = ['', 'nya', 'ku', 'mu', 'lah', 'kah', 'pun', 'kan', 'i', 'an', 'annya', 'kannya', 'innya']
    
    # Kupas Akhiran Dulu
    for s in suffixes:
        if s and kata.endswith(s):
            if kata[:-len(s)] in INDO_DB: return True
            
    # Kupas Awalan + Akhiran
    for p in prefixes:
        if p and kata.startswith(p):
            roots = [kata[len(p):]]
            if p in ['meny', 'peny']: roots.append('s' + kata[len(p):])
            elif p in ['meng', 'peng']: roots.extend(['k' + kata[len(p):], 'g' + kata[len(p):], 'h' + kata[len(p):]])
            elif p in ['mem', 'pem']: roots.append('p' + kata[len(p):])
            elif p in ['men', 'pen']: roots.append('t' + kata[len(p):])
            
            for root in roots:
                if not root: continue
                if root in INDO_DB: return True
                for s in suffixes:
                    if s and root.endswith(s):
                        if root[:-len(s)] in INDO_DB:
                            return True
    return False

def is_valid_eng(kata):
    """Cek English murni + Akhiran khas"""
    if not kata: return False
    
    # Penanganan kata ber-strip (contoh: closed-loop)
    if '-' in kata:
        return all(is_valid_eng(p) for p in kata.split('-') if p)
        
    if kata in ENG_DB: return True
    
    suffixes = ['s', 'es', 'ed', 'ing', 'ly', 'less', 'ness', 'ment']
    for suf in suffixes:
        if kata.endswith(suf):
            root = kata[:-len(suf)]
            if root in ENG_DB: return True
            if root + 'e' in ENG_DB: return True
            if len(root) > 1 and root[-1] == root[-2]: 
                if root[:-1] in ENG_DB: return True
                
    patterns = [r'tion$', r'ity$', r'ous$', r'ish$', r'q', r'x', r'tch', r'ph']
    if any(re.search(p, kata) for p in patterns): return True
    
    return False

def get_canonical(word):
    return word.lower().replace('-', '')

def bikin_regex_kebal(kata):
    kata_bersih = re.sub(r'[\s\-\u2010-\u2015_]+', '', kata.strip())
    if not kata_bersih: return ""
    return r'\b' + r'[\s\-\u2010-\u2015_]*'.join([re.escape(h) for h in kata_bersih]) + r'\b'

def get_semua_paragraf(doc):
    paragraf_list = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    paragraf_list.append(p)
    return paragraf_list

def proses_italic_aman(doc, daftar_kata):
    # Urutkan kata dari yang terpanjang ke yang terpendek agar frasa seperti "data logger" 
    # dieksekusi duluan sebelum kata "logger" atau "data" tunggal.
    daftar_kata = sorted([k.strip() for k in daftar_kata if k.strip()], key=len, reverse=True)
    
    patterns = [bikin_regex_kebal(kata) for kata in daftar_kata]
    if not patterns: return doc
    regex_master = re.compile(f"({'|'.join(patterns)})", re.IGNORECASE)
    
    for p in get_semua_paragraf(doc):
        if not p.text.strip(): continue
        matches = list(regex_master.finditer(p.text))
        if not matches: continue
        
        char_styles = [0] * len(p.text)
        for m in matches:
            for i in range(m.start(), m.end()): char_styles[i] = 1
                
        run_starts = []
        curr = 0
        for run in p.runs:
            run_starts.append(curr)
            curr += len(run.text) if run.text else 0
            
        runs = p.runs
        for i in range(len(runs)-1, -1, -1):
            run = runs[i]
            if not run.text: continue
            
            start_idx = run_starts[i]
            styles = char_styles[start_idx : start_idx + len(run.text)]
            if 1 not in styles: continue
            if 0 not in styles:
                run.italic = True
                continue
                
            parts = []
            cur_style, cur_text = styles[0], run.text[0]
            for j in range(1, len(run.text)):
                if styles[j] == cur_style: cur_text += run.text[j]
                else:
                    parts.append((cur_text, cur_style))
                    cur_style, cur_text = styles[j], run.text[j]
            parts.append((cur_text, cur_style))
            
            run.text = parts[0][0]
            if parts[0][1] == 1: run.italic = True
            
            cur_r = run._r
            for text, style in parts[1:]:
                new_r = OxmlElement('w:r')
                cur_r.addnext(new_r)
                if run._r.rPr is not None: new_r.append(deepcopy(run._r.rPr))
                new_run = type(run)(new_r, run._parent)
                new_run.text = text
                if style == 1: new_run.italic = True
                cur_r = new_r
    return doc


# ─── ENDPOINTS ───
@app.get("/")
def home():
    return {"status": "API Offline Paten (Word-by-Word Scanner) Ready!"}

@app.post("/api/scan")
async def scan_document(file: UploadFile = File(...)):
    doc = Document(io.BytesIO(await file.read()))
    full_text = " ".join([p.text for p in get_semua_paragraf(doc) if p.text.strip()])
    
    eng_dict = {}
    unk_dict = {}
    
    # 1. Menangkap semua kata, termasuk yang memiliki tanda strip (-)
    words_in_doc = re.finditer(r'\b[a-zA-Z\-]+\b', full_text)
    
    word_cases = {}
    for match in words_in_doc:
        w = match.group(0)
        wl = w.lower()
        if wl not in word_cases: word_cases[wl] = []
        word_cases[wl].append(w)
        
    # 2. Logika Utama: Cek Kata per Kata
    for w_lower, originals in word_cases.items():
        # Abaikan kata 1-2 huruf
        if len(w_lower) <= 2: continue 
        
        # Abaikan jika selalu ditulis huruf Kapital di dokumen (Kemungkinan Nama / Singkatan)
        is_always_cap = all(orig.istitle() or orig.isupper() for orig in originals)
        
        # ── KONDISI 1: APAKAH INI BAHASA INDONESIA? ──
        if is_valid_indo(w_lower):
            continue # Abaikan! Kata ini (spt "data", "relay", "memperoleh") aman.
            
        # Jika bukan Indo, tapi selalu Kapital, buang (Pasti Nama Orang, spt MOCO, dll)
        if is_always_cap:
            continue
            
        # ── KONDISI 2: APAKAH INI BAHASA INGGRIS? ──
        if is_valid_eng(w_lower):
            canon = get_canonical(w_lower)
            if canon not in eng_dict or len(w_lower) < len(eng_dict[canon]):
                eng_dict[canon] = w_lower
        
        # ── KONDISI 3: BUKAN KEDUANYA (TYPO / ASING LAIN) ──
        else:
            canon = get_canonical(w_lower)
            if canon not in unk_dict or len(w_lower) < len(unk_dict[canon]):
                unk_dict[canon] = w_lower
                
    return {
        "english_terms": sorted(list(eng_dict.values())),
        "unknown_terms": sorted(list(unk_dict.values()))
    }

@app.post("/api/process")
async def process_document(file: UploadFile = File(...), terms: str = Form(...)):
    kata_list = [k.strip() for k in terms.split('\n') if k.strip()]
    doc_hasil = proses_italic_aman(Document(io.BytesIO(await file.read())), kata_list)
    output_stream = io.BytesIO()
    doc_hasil.save(output_stream)
    output_stream.seek(0)
    return StreamingResponse(
        output_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Italify_{file.filename}"'}
    )

@app.post("/api/preview")
async def preview_document(file: UploadFile = File(...), terms: str = Form(default="")):
    doc = Document(io.BytesIO(await file.read()))
    kata_list = [k.strip() for k in terms.split('\n') if k.strip()]
    
    # Sortir juga buat preview biar akurat
    kata_list = sorted(kata_list, key=len, reverse=True)
    patterns_preview = [bikin_regex_kebal(kata) for kata in kata_list if kata.strip()]
    
    match_count, preview_html = 0, ""
    
    if patterns_preview: regex_preview = re.compile(f"({'|'.join(patterns_preview)})", re.IGNORECASE)
        
    for p in get_semua_paragraf(doc):
        teks = p.text.strip()
        if not teks: continue
        teks_esc = teks.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if patterns_preview:
            teks_hl, n_subs = regex_preview.subn(lambda m: f"<i>{m.group(0)}</i>", teks_esc)
            match_count += n_subs
            preview_html += f'<p>{teks_hl}</p>'
        else:
            preview_html += f'<p>{teks_esc}</p>'
            
    if not preview_html: preview_html = "<p style='text-align:center; margin-top:20%; color:var(--text-muted);'>Kosong</p>"
    return {"html": preview_html, "match_count": match_count}